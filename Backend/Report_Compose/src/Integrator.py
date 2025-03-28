import asyncio
import json
import logging
import os
import pickle
import random
import tempfile
import subprocess
import urllib.parse


import yfinance as yf
from alpha_vantage.timeseries import TimeSeries
from alpha_vantage.fundamentaldata import FundamentalData
from alpha_vantage.alphavantage import AlphaVantage

from typing import Dict
from xml.etree.ElementTree import indent

import networkx as nx
import yaml

from Backend.Report_Compose.src.DataMolder import DataMolder
from Backend.Report_Compose.src.DataQuerier import DataQuerier
from Backend.Report_Compose.src.ResultsDAG import ResultsDAG
from Backend.Report_Compose.src.PromptManager import PromptManager
import json
import httpx


def load_api_key(yaml_file_path: str, api_name: str) -> str:
    """
    Loads the OpenAI API key from the given YAML file.
    Implements error handling for missing file, parsing errors, or missing key.
    """
    try:
        # Check if file exists
        if not os.path.exists(yaml_file_path):
            raise FileNotFoundError(f"YAML file not found: {yaml_file_path}")

        # Open and parse the YAML file
        with open(yaml_file_path, "r") as file:
            data = yaml.safe_load(file)

        # Validate if the OpenAI key exists
        if "API_Keys" not in data or "OpenAI" not in data["API_Keys"]:
            raise KeyError("OpenAI API key is missing from the YAML file.")

        return data["API_Keys"][api_name]

    except FileNotFoundError as e:
        print(f"Error: {e}")
    except yaml.YAMLError as e:
        print(f"Error parsing YAML file: {e}")
    except KeyError as e:
        print(f"Error: {e}")
    return ""



class Integrator:
    """
    This class serves as a orchestrator of the services necessary to compose the report based on the YAML file.
    Integrator interacts with Prompt_Manager to store and retrieve the prompts from a YAML file.

    Although composing the report could all be done in one shot, since the time it would take to complete the report
    by the LLM on the back-end are at the moment too long. It is important that each completed prompt node
    is returned as we complete the prompt. The idea is that Integrator interacts with RequestMngrAPI who is
    responsible for handling the incoming report generation request that are coming from the front-end.

    The YAML prompt file contains the prompts (section_name, text, id, system) and the prompt_dag which corresponds
    to a graph indicating how prompts depend on each other. Although PromptManager takes care of the Directed-Acyclic
    Graph, the prompts need to be explored by Integrator in topological order such that the parents prompt results
    are added to the context of the corresponding child prompts.

    The dag needs to be explored asynchronously per each node, and once any non-system (system = false)
     node DataQuery data is retrieved, then such output needs to be sent to DataMolder which will accept all the
     previous context and responses (in consecutive order) plus the results of the Web Search done by the corresponding
     Data Querier node.

     The result output from Data Molder (which is a LLM-based program) then becomes the current node's output.

     As the Directed-Acyclic-Graph gets completed; the results for each node should be placed on DAG structure
     in our case we will have a separated file: ResultsDAG.py which gets updated every time a node is completed.

    A great approach would be to BFS the prompts inside PromptManager and every time some node gets completed; store it
    in a node inside the ResultsDAG object.

    Ultimately, the ResultsDAG object should be accessible through our front-end-facing class "RequestMngrAPI.py"
    that will be the instantiator of all the process. Such API will have a WebSocket with the front-end (client)
    so that the client PC is able to explore the resulting information.


    """

    def __init__(self, yaml_file_path: str):
        """
        Initialize the Integrator with a path to the prompts YAML.
        Creates a PromptManager and a fresh ResultsDAG.
        """
        self.molder_model = "gpt-4o"
        self.yaml_file_path = yaml_file_path  # <-- Add this line to store the file path
        self.prompt_manager = PromptManager(yaml_file_path)
        self.results_dag = ResultsDAG()
        self.tasks = {}
        self.openAI_API_key = load_api_key("./Credentials/Credentials.yaml", "OpenAI")
        self.polygon_api_key = load_api_key("./Credentials/Credentials.yaml", "Polygon")
        self.alpha_vantage_api_key = load_api_key("./Credentials/Credentials.yaml", "Vantage")
        self.focus_message = "Default Focus Message"
        self.fin_numeric_cntxt = {"default": 0}
        self.web_search = True
        self.failed_nodes = asyncio.Queue()
        self.node_attempts = {}

        # At the end of the __init__ method of the Integrator class, add the following method:

    def __getstate__(self):
        state = self.__dict__.copy()
        if 'tasks' in state:
            logging.info("Removing 'tasks' attribute from Integrator state for pickling")
            del state['tasks']
        # Log the picklability of each attribute
        for key, value in state.items():
            try:
                pickle.dumps(value)
                logging.info("Attribute '%s' is pickleable", key)
            except Exception as ex:
                logging.error("Attribute '%s' is NOT pickleable: %s", key, ex)
        return state

    import asyncio
    import logging
    import httpx

    async def get_search_api_url(self):
        """
        Concurrently checks candidate base URLs by polling their /health endpoint every 10 seconds,
        for up to 2 minutes. Returns the /search endpoint of the first candidate to respond successfully.
        Raises an exception if none respond within the timeout.
        """
        candidates = ["http://localhost:8383", "http://web_search_api:8383"]

        async def check_candidate(base_url: str) -> str:
            health_url = f"{base_url}/health"
            timeout = 1000  # seconds for the entire candidate
            interval = 10  # poll every 10 seconds
            end_time = asyncio.get_running_loop().time() + timeout
            async with httpx.AsyncClient() as client:
                while asyncio.get_running_loop().time() < end_time:
                    try:
                        response = await client.get(health_url, timeout=15.0)
                        if response.status_code == 200 and response.json().get("status") == "ok":
                            return f"{base_url}/search"
                    except Exception as e:
                        logging.info("Health check failed for %s: %s", base_url, e)
                    await asyncio.sleep(interval)
            return None

        tasks = [asyncio.create_task(check_candidate(url)) for url in candidates]
        done, pending = await asyncio.wait(tasks, return_when=asyncio.FIRST_COMPLETED)

        # Cancel any pending tasks
        for task in pending:
            task.cancel()

        # Check if any candidate succeeded
        for task in done:
            result = task.result()
            if result:
                return result

        raise Exception("No available Search API endpoint after health checks.")

    async def process_node(self, node_id: int, focus_message) -> tuple[None, None] | tuple[str, any]:
        curr_prompt = self.prompt_manager.get_prompt_by_id(node_id)

        if curr_prompt['system'] is True:
            logging.info(f"Skipping node {node_id} since it's system prompt")
            return "**This is a system prompt**", {"results": [{"System Node": "NA_system_node"}]}

        if self.web_search:
            # Dynamically determine the Search API URL.
            # TODO: The search_api_url may return an exception of failure if
            # the api endpoint is to busy and won't send a response back. Therefore,
            # we must try to get the search_api_url every 5 seconds as longs as there
            # are other nodes being processed. I still do not know how to properly code this
            # since inside this scope there is no way to know if there are nodes being executed (I think).
            search_api_url = await self.get_search_api_url()
            ###
            querier = DataQuerier(curr_prompt['text'], focus_message, search_api_url)
            print(f'Processing node {node_id} with prompt: {json.dumps(curr_prompt, indent=4)}')
            await querier.query_and_process()
            online_data = querier.get_processed_data()
        else:
            online_data = {"results": [{"mock_data": "place_holder"}]}
        print(f'Count of articles found for node {node_id}: {len(online_data)}')
        self.molder_model = "gpt-4o-search-preview"
        molder = DataMolder(self.molder_model, self.openAI_API_key)
        ancestor_messages = self.get_ancestor_chat_hist(node_id).copy()
        logging.info("Dumping Numerical Context")
        num_context_message = "Here is some data for context" + json.dumps(self.fin_numeric_cntxt, indent=4)
        logging.info("Dumped Numerical Context")
        numeric_context = {"text": num_context_message, "entity": "user"}
        ancestor_messages.insert(1, numeric_context)
        if self.molder_model == "gpt-4o-search-preview":
            if len(ancestor_messages) <= 2:
                # If the list is too short, use all of it.
                ancestor_messages = ancestor_messages
            else:
                # Otherwise, take the first two and the last one.
                ancestor_messages = ancestor_messages[:2] + [ancestor_messages[-1]]
                logging.info(f'Ancestor count of node {node_id}: {len(ancestor_messages)}')
        molded_tup = await molder.process_data(online_data, ancestor_messages, focus_message)
        response = molded_tup["llm_response"]
        llm_found_online_data = molded_tup["web_references"]

        # Parse the newline-separated web references and create dictionary entries.
        if "results" in online_data and isinstance(online_data["results"], list):
            new_refs = []
            cached_urls = []
            for line in llm_found_online_data.splitlines():
                line = line.strip()
                if not line:
                    continue
                # Each line is expected to be in the format "Title: URL"
                parts = line.split(":", 1)
                if len(parts) == 2:
                    title = parts[0].strip()
                    url = parts[1].strip()
                else:
                    title = line
                    url = ""
                if url in cached_urls:
                    continue
                # Create a dictionary in the same format as online_data items.
                cached_urls.append(url)
                ref_dict = {
                    "url": url,
                    "display_url": url,  # For simplicity, using the full URL.
                    "snippet": "",
                    "title": title,
                    "scrapped_text": "",
                    "extension": "html"
                }
                new_refs.append(ref_dict)
            # Prepend the new web reference entries to the existing results.
            online_data["results"] = new_refs + online_data["results"]
        else:
            logging.info(f"No results found for node {node_id}")

        return response, online_data

    def get_ancestor_chat_hist(self, node_id: int) -> list[Dict[str, any]]:
        """
        This method is necessary for building the chat history so far at a current node
        so only the relevant chat history is later sent to be processed by the LLM.
        How it works is that given the results_dag data-structure; the method references
        the correct node by ID then backtraces to root node.
        Then from 0 (current node) it retrieves both the response of the node and the
        prompt for such node. Is important to note that the retrieval is done by querying
        the PromptManager and ResultsDAG correspondingly.
        Since the ResultsDAG contains responses from the LLM, while the PromptManager contains the
        messages sent to the LLM (from user or system), then this is necessary to inform leveraging
        the Dictionary. Some nodes may not contain an LLM response from the RAG.
        :param node_id: ID of the node to retrieve the chat history for.
        :return: A list of the chat history for the given node.
                Example:
                    [{entity: 'system',
                      text: 'system_prompt',
                      },
                     {entity: 'user',
                      text: 'user prompt',
                    },
                     {entity: 'llm',
                     text: 'llm response'
                     }
                    ]
                The system prompts will not be followed by responses from the LLM.
        """
        dag = self.prompt_manager.prompt_dag

        # 1. Collect all ancestors of 'node_id', plus the node itself
        ancestors = nx.ancestors(dag, node_id)

        relevant_nodes = ancestors.union({node_id})

        # 2. Perform a topological sort of the entire DAG
        full_order = list(nx.topological_sort(dag))

        # 3. Filter to only those that lead to 'node_id'
        path_nodes = [n for n in full_order if n in relevant_nodes]

        chat_history = []

        for n in path_nodes:
            prompt = self.prompt_manager.get_prompt_by_id(n)
            if not prompt:
                # Should never happen if the YAML is valid, but just in case
                continue

            # 4a. Add the node's own prompt (system => "system", else => "user")
            if prompt["system"]:
                chat_history.append({
                    "entity": "system",
                    "text": prompt["text"]
                })
            else:
                chat_history.append({
                    "entity": "user",
                    "text": prompt["text"]
                })

            # 4b. If there's a completed LLM response, add it (optional business logic)
            node_result = self.results_dag.get_result(n)
            if node_result and node_result["status"] == "complete":
                # By default, we only add LLM response for non-system prompts,
                # but adjust if your logic differs.
                if not prompt["system"]:
                    chat_history.append({
                        "entity": "llm",
                        "text": str(node_result["result"]['llm'])  # Using llm here is very important to avoid
                        # input overflow of ancestor online_data
                    })

        return chat_history

    async def queue_node(self, node_id: int, dag: nx.DiGraph, mock: bool):
        """
        Process a single node in the DAG. This task will:
          1. Wait for all parent node tasks to complete.
          2. Run the node's processing (mock or real).
          3. Store the result or mark failure.
        """
        parent_ids = list(dag.predecessors(node_id))
        if parent_ids:
            await asyncio.gather(*(self.tasks[parent_id] for parent_id in parent_ids))

        try:
            # <--- 1) Mark the node as processing right here
            self.results_dag.mark_processing(node_id, "Node is currently being explored")
            result = {}
            node_prompt = self.prompt_manager.get_prompt_by_id(node_id)
            node_name = node_prompt["section_title"]
            if mock:
                # Simulate processing
                process_time = abs(random.gauss(5, 2))
                await asyncio.sleep(process_time)
                result = {'llm': "Some llm response", "online_data": "some_online_data"}
            else:
                # IMPORTANT: Await the async call to process_node so that you store the final result.
                response, online_data = await self.process_node(node_id, self.focus_message)
                result = {'llm': response, 'online_data': online_data}
                # logging.info(f'Result node {node_id}: {result}')
            result['section_tile'] = node_name
            self.results_dag.store_result(node_id, result)
        except Exception as e:
            self.results_dag.mark_failed(node_id, str(e))

    ######################################
    ######################################
    ######################################
    #####################################
    ####################
    ###################
    ###################
    ######################################
    ######################################
    ###################

    async def generate_report(self, focus_message: str, mock: bool = False, web_search: bool = True,
                              company: bool = True) -> str:
        """
        Process the prompt DAG concurrently.
        Each node is scheduled as soon as its dependencies are complete.
        The web_search flag is set based on the API parameter.
        """
        if company:
            from Backend.Report_Compose.src.FinancialDataRetriever import FinancialDataRetriever
            retriever = FinancialDataRetriever(
                alpha_vantage_api_key=self.alpha_vantage_api_key,
                polygon_api_key=self.polygon_api_key
            )
            self.fin_numeric_cntxt = await retriever.get_financial_info_yahoo(focus_message)

        self.focus_message = focus_message
        self.web_search = web_search  # Propagate the parameter to the integrator

        dag = self.prompt_manager.prompt_dag

        # Initialize each node in ResultsDAG as "pending"
        for node_id in dag.nodes():
            self.results_dag.init_node(node_id)

        # Schedule tasks for each node in topological order
        for node_id in nx.topological_sort(dag):
            self.tasks[node_id] = asyncio.create_task(self.queue_node(node_id, dag, mock))

        # Await all node tasks concurrently
        await asyncio.gather(*self.tasks.values())

        #print(json.dumps(self.results_dag.to_json(), indent=4))
        return self.results_dag.to_json()

    def generate_docx_report(self, llm_format: str = "Markdown") -> str:
        """
        Generates a DOCX report using the final DAG results.

        Parameters:
            llm_format (str): The format of the LLM response.
                              Currently, only "Markdown" is supported.
                              Future formats (e.g. LaTeX) can be implemented here.

        Process:
          1. Parse the final DAG results (stored as a JSON string) into a Python dictionary.
          2. Obtain the node order via a topological sort from the prompt DAG.
          3. Extract the prompt set name from the YAML file path.
          4. Delegate to a helper method based on the specified llm_format.

        Returns:
            str: The temporary file path of the generated DOCX report.

        Critical Decisions:
          - We assume the LLM responses are in Markdown format.
          - By delegating to a helper method, we keep the report-generation logic modular.
        """
        # Parse the results DAG.
        try:
            dag_obj = json.loads(self.results_dag.to_json())
        except Exception as e:
            raise Exception(f"Error parsing DAG data: {e}")

        # Get the node order from the prompt DAG.
        dag_graph = self.prompt_manager.prompt_dag
        node_order = list(nx.topological_sort(dag_graph))

        # Extract the prompt set name.
        prompt_set = os.path.basename(self.yaml_file_path).replace(".yaml", "")

        # Delegate to the helper method based on llm_format.
        if llm_format.lower() == "markdown":
            return self.generate_docx_from_md(dag_obj, node_order, prompt_set)
        else:
            # Future implementations could support additional formats.
            raise Exception(f"Unsupported llm_format: {llm_format}. Only 'Markdown' is currently supported.")

    def generate_docx_from_md(self, dag_obj: dict, node_order: list, prompt_set: str) -> str:
        """
        Helper method that generates a DOCX report from Markdown-formatted LLM responses.

        Report Structure:
          - Title and Metadata: The report title, prompt set, and focus message.
          - Main Sections: For each node, add a section heading (with the node's section title)
            and insert the LLM response as plain text.
          - Appendix: A dedicated section listing the online data for each node.

        Parameters:
            dag_obj (dict): The parsed results DAG.
            node_order (list): A topologically sorted list of node IDs.
            prompt_set (str): The prompt set name, derived from the YAML file path.

        Returns:
            str: The temporary file path of the generated DOCX report.

        Critical Decisions:
          - We assume that the LLM responses are in Markdown. For now, we insert the Markdown
            as plain text into the document. In the future, a richer conversion (e.g., via pypandoc)
            could be implemented.
          - The report is split into main sections (for each node) and an appendix for online data.
        """
        from docx import Document

        # Create a new document.
        doc = Document()
        doc.add_heading("Aggregated Report", 0)
        doc.add_paragraph(f"Prompt Set: {prompt_set}")
        doc.add_paragraph(f"Focus Message: {self.focus_message}")
        doc.add_paragraph("")  # Blank line for spacing.

        # Main report sections.
        for idx, node_id in enumerate(node_order, start=1):
            node_result = dag_obj.get(str(node_id)) or dag_obj.get(node_id)
            if not node_result:
                continue
            # Extract the section title from the node result.
            section_title = (node_result.get("result", {}).get("section_tile") or
                             node_result.get("result", {}).get("section_title") or
                             f"Section {idx}")
            # Extract the LLM response.
            llm_response = node_result.get("result", {}).get("llm", "No LLM response found.")

            doc.add_heading(f"Section {idx}. {section_title}", level=1)
            doc.add_paragraph(llm_response)

        # Appendix for online data.
        """doc.add_heading("Appendix: Online Data", level=1)
        for idx, node_id in enumerate(node_order, start=1):
            node_result = dag_obj.get(str(node_id)) or dag_obj.get(node_id)
            if not node_result:
                continue
            online_data = node_result.get("result", {}).get("online_data", "No online data found.")
            section_title = (node_result.get("result", {}).get("section_tile") or
                             node_result.get("result", {}).get("section_title") or
                             f"Section {idx}")
            doc.add_heading(f"Section {idx} Online Data - {section_title}", level=2)
            doc.add_paragraph(str(online_data))
        """
        # ------------------------------
        # Old Appendix for online data (to be replaced)
        # ------------------------------
        # doc.add_heading("Appendix: Online Data", level=1)
        # for idx, node_id in enumerate(node_order, start=1):
        #     node_result = dag_obj.get(str(node_id)) or dag_obj.get(node_id)
        #     if not node_result:
        #         continue
        #     online_data = node_result.get("result", {}).get("online_data", "No online data found.")
        #     section_title = (node_result.get("result", {}).get("section_tile") or
        #                      node_result.get("result", {}).get("section_title") or
        #                      f"Section {idx}")
        #     doc.add_heading(f"Section {idx} Online Data - {section_title}", level=2)
        #     doc.add_paragraph(str(online_data))

        # ------------------------------
        # New References Section for Online Data
        # ------------------------------
        # Import necessary OXML helpers for hyperlink creation.
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        import docx.opc.constants

        # Helper function to add a hyperlink to a paragraph.
        def add_hyperlink(paragraph, url, text, color="0000FF", underline=True):
            # This function creates a hyperlink within a paragraph.
            part = paragraph.part
            r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
            hyperlink = OxmlElement('w:hyperlink')
            hyperlink.set(qn('r:id'), r_id)
            new_run = OxmlElement('w:r')
            rPr = OxmlElement('w:rPr')
            # Set color.
            c = OxmlElement('w:color')
            c.set(qn('w:val'), color)
            rPr.append(c)
            # Set underline property.
            u = OxmlElement('w:u')
            u.set(qn('w:val'), 'single' if underline else 'none')
            rPr.append(u)
            new_run.append(rPr)
            new_run.text = text
            hyperlink.append(new_run)
            paragraph._p.append(hyperlink)
            return hyperlink

        doc.add_heading("References", level=1)
        # For each node, if online data is available, list its references.
        for idx, node_id in enumerate(node_order, start=1):

            node_result = dag_obj.get(str(node_id)) or dag_obj.get(node_id)
            if not node_result:
                continue
            online_data = node_result.get("result", {}).get("online_data", None)
            if not online_data or "results" not in online_data:
                continue
            section_title = (node_result.get("result", {}).get("section_tile") or
                             node_result.get("result", {}).get("section_title") or
                             f"Section {idx}")
            doc.add_heading(f"References for {section_title}", level=2)
            for res in online_data["results"]:
                # Create a reference box.
                box_para = doc.add_paragraph()
                # Add the title as hyperlinked bold text if URL is present.
                if res.get("title") and res.get("url"):
                    run = box_para.add_run("")
                    add_hyperlink(box_para, res.get("url"), res.get("title"), color="0000FF", underline=True)
                else:
                    box_para.add_run(res.get("title", "No Title")).bold = True

                # Add the scrapped text below, if available.
                if res.get("scrapped_text"):
                    doc.add_paragraph(res.get("scrapped_text"), style="Intense Quote")
                # Add display_url as a smaller source line, if available.
                if res.get("display_url") and res.get("url"):
                    src_para = doc.add_paragraph("Source: ")
                    add_hyperlink(src_para, res.get("url"), res.get("display_url"), color="808080", underline=False)
                # Add a separator.
                doc.add_paragraph("------------------------------")

        # Save the document to a temporary file.
        tmp_file = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
        doc.save(tmp_file.name)
        tmp_file.close()
        return tmp_file.name

    def generate_pdf_report(self) -> str:
        """
        Generates a PDF report using the final DAG results by first creating a DOCX report,
        then converting it to PDF using LibreOffice in headless mode.
        Returns the temporary file path of the generated PDF.
        """
        # First, generate the DOCX report.
        docx_path = self.generate_docx_report()

        # Create a temporary directory for the PDF output.
        output_dir = tempfile.mkdtemp()

        # Construct the LibreOffice command to convert DOCX to PDF.
        command = [
            "libreoffice",
            "--headless",
            "--convert-to", "pdf",
            docx_path,
            "--outdir", output_dir
        ]

        try:
            subprocess.run(command, check=True, stdout=subprocess.PIPE, stderr=subprocess.PIPE)
        except subprocess.CalledProcessError as e:
            raise Exception(f"Error converting DOCX to PDF: {e.stderr.decode('utf-8')}")

        # The converted file will have the same basename as the DOCX but with .pdf extension.
        pdf_file = os.path.join(output_dir, os.path.basename(docx_path).replace(".docx", ".pdf"))
        if not os.path.exists(pdf_file):
            raise Exception("PDF conversion failed; output file not found.")

        return pdf_file



